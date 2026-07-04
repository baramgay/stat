"""Output Window — 독립 결과 창 (탭 분리 + 클립보드 복사 + Word/HTML 내보내기).

분석 결과별로 탭을 생성하고, 우클릭으로 표를 클립보드에 복사하거나
Word/HTML로 내보낼 수 있습니다.
"""

from __future__ import annotations

import html as html_mod
import io
from datetime import datetime
from typing import Any

from PySide6.QtCore import Qt
from PySide6.QtGui import QAction, QKeySequence, QTextCursor
from PySide6.QtWidgets import (
    QApplication,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QMainWindow,
    QMenu,
    QMessageBox,
    QPushButton,
    QTabWidget,
    QTextBrowser,
    QVBoxLayout,
    QWidget,
)

from nuristat.ui.theme import get_output_html_styles

# 차트 편집기 — 선택 의존 (PySide6 설치 환경에서만 활성)
try:
    from nuristat.ui.dialogs.chart_editor_dialog import ChartEditorDialog as _ChartEditorDialog
    _CHART_EDITOR_AVAILABLE = True
except Exception:
    _CHART_EDITOR_AVAILABLE = False


class OutputWindow(QMainWindow):
    """독립 결과 창.

    - 분석 결과: 탭 1개 = 분석 1건 (``add_analysis_result``)
    - 로그/상태: "로그" 탭 누적 (``add_output``)
    - 우클릭 → 클립보드 복사 (text/html + text/plain 탭구분)
    - 파일 메뉴 → HTML / Word 내보내기
    """

    def __init__(self, parent=None) -> None:
        super().__init__(parent)
        self.setWindowTitle("📊 누리스탯 결과")
        self.setMinimumSize(700, 500)
        self.resize(900, 650)

        self._results: list[Any] = []   # AnalysisResult per tab (aligned with tabs 1..)
        self._log_lines: list[str] = []

        self._setup_ui()
        self._setup_menus()

    # ------------------------------------------------------------------
    # UI 초기화
    # ------------------------------------------------------------------

    def _setup_ui(self) -> None:
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(6, 6, 6, 6)
        layout.setSpacing(4)

        # 상단 버튼 바
        bar = QHBoxLayout()
        self.info_label = QLabel("분석 결과가 탭으로 표시됩니다")
        self.info_label.setStyleSheet("color: #55555f; font-size: 11px;")
        bar.addWidget(self.info_label)
        bar.addStretch()

        self.btn_clear = QPushButton("🗑️ 지우기")
        self.btn_clear.setToolTip("모든 탭을 지웁니다")
        self.btn_clear.clicked.connect(self.clear_output)
        bar.addWidget(self.btn_clear)

        self.btn_save_html = QPushButton("💾 HTML")
        self.btn_save_html.setToolTip("현재 탭을 HTML로 저장합니다")
        self.btn_save_html.clicked.connect(self._save_html)
        bar.addWidget(self.btn_save_html)

        self.btn_save_word = QPushButton("📄 Word")
        self.btn_save_word.setToolTip("현재 탭을 Word(.docx)로 저장합니다")
        self.btn_save_word.clicked.connect(self._save_word)
        bar.addWidget(self.btn_save_word)

        self.btn_edit_chart = QPushButton("✏️ 차트 수정")
        self.btn_edit_chart.setToolTip("현재 탭의 차트 텍스트/레이아웃을 편집합니다")
        self.btn_edit_chart.clicked.connect(self._edit_chart)
        bar.addWidget(self.btn_edit_chart)

        layout.addLayout(bar)

        # 탭 위젯
        self.tab_widget = QTabWidget()
        self.tab_widget.setTabsClosable(True)
        self.tab_widget.tabCloseRequested.connect(self._on_tab_close)
        layout.addWidget(self.tab_widget)

        # 로그 탭 (항상 존재, index 0)
        self._log_browser = self._make_browser()
        self.tab_widget.addTab(self._log_browser, "📋 로그")
        from PySide6.QtWidgets import QTabBar
        self.tab_widget.tabBar().setTabButton(
            0, QTabBar.ButtonPosition.RightSide, None
        )

        self.statusbar = self.statusBar()
        self.statusbar.showMessage("준비됨")

    def _make_browser(self) -> QTextBrowser:
        """새 QTextBrowser를 생성하고 컨텍스트 메뉴를 연결합니다."""
        browser = QTextBrowser()
        browser.setOpenLinks(False)
        browser.setContextMenuPolicy(Qt.CustomContextMenu)
        browser.customContextMenuRequested.connect(
            lambda pos, b=browser: self._show_context_menu(pos, b)
        )
        return browser

    def _setup_menus(self) -> None:
        menubar = self.menuBar()

        file_menu = menubar.addMenu("파일(&F)")

        save_html_action = QAction("HTML로 저장(&S)", self)
        save_html_action.setShortcut(QKeySequence.Save)
        save_html_action.triggered.connect(self._save_html)
        file_menu.addAction(save_html_action)

        save_word_action = QAction("Word(.docx)로 저장(&W)", self)
        save_word_action.triggered.connect(self._save_word)
        file_menu.addAction(save_word_action)

        file_menu.addSeparator()

        close_action = QAction("닫기(&C)", self)
        close_action.setShortcut(QKeySequence.Close)
        close_action.triggered.connect(self.hide)
        file_menu.addAction(close_action)

        edit_menu = menubar.addMenu("편집(&E)")
        clear_action = QAction("모두 지우기", self)
        clear_action.triggered.connect(self.clear_output)
        edit_menu.addAction(clear_action)

        edit_menu.addSeparator()
        chart_edit_action = QAction("차트 수정(&E)", self)
        chart_edit_action.triggered.connect(self._edit_chart)
        edit_menu.addAction(chart_edit_action)

    # ------------------------------------------------------------------
    # 공개 API
    # ------------------------------------------------------------------

    def add_analysis_result(self, result: Any) -> None:
        """분석 결과 탭 추가.

        result는 AnalysisResult 또는 to_html() 메서드를 가진 객체여야 합니다.
        """
        ts = datetime.now().strftime("%H:%M:%S")
        title = getattr(result, "title", "결과")
        tab_label = f"{title[:12]} {ts}"

        browser = self._make_browser()
        html = self._wrap_html(result.to_html(), title, ts)
        browser.setHtml(html)

        tab_idx = self.tab_widget.addTab(browser, tab_label)
        self.tab_widget.setCurrentIndex(tab_idx)

        self._results.append(result)
        self.info_label.setText(f"분석 결과 {len(self._results)}건")
        self.statusbar.showMessage(f"[{ts}] {title} 완료")

    def add_output(self, content: str, output_type: str = "text") -> None:
        """로그/상태 메시지를 '로그' 탭에 누적합니다.

        기존 코드와 호환성을 유지합니다 (success/error/warning/analysis).
        """
        ts = datetime.now().strftime("%H:%M:%S")
        color_map = {
            "success": "#27ae60",
            "error":   "#d62728",
            "warning": "#ff7f0e",
            "analysis": "#2874a6",
            "text":    "#333333",
        }
        color = color_map.get(output_type, "#333333")
        line = (
            f'<div style="margin:4px 0;">'
            f'<span style="color:#55555f; font-size:10px;">{ts}</span> '
            f'<span style="color:{color}">{html_mod.escape(content)}</span>'
            f'</div>'
        )
        self._log_lines.append(line)
        cursor = self._log_browser.textCursor()
        cursor.movePosition(QTextCursor.End)
        cursor.insertHtml(line)
        sb = self._log_browser.verticalScrollBar()
        sb.setValue(sb.maximum())
        self.statusbar.showMessage(f"로그: {len(self._log_lines)}건")

    def clear_output(self) -> None:
        """분석 탭(인덱스 1+)을 모두 닫고 로그도 지웁니다."""
        while self.tab_widget.count() > 1:
            self.tab_widget.removeTab(1)
        self._results.clear()
        self._log_lines.clear()
        self._log_browser.clear()
        self.info_label.setText("분석 결과가 탭으로 표시됩니다")
        self.statusbar.showMessage("출력이 지워졌습니다")

    # ------------------------------------------------------------------
    # 컨텍스트 메뉴 (우클릭 → 클립보드 복사)
    # ------------------------------------------------------------------

    def _show_context_menu(self, pos, browser: QTextBrowser) -> None:
        result = self._result_for_browser(browser)
        menu = QMenu(self)

        if result is not None:
            copy_html_act = QAction("📋 표 복사 (한글/Word 붙여넣기)", self)
            copy_html_act.triggered.connect(lambda: self._copy_tables(result))
            menu.addAction(copy_html_act)

            copy_text_act = QAction("📄 텍스트 복사 (탭 구분)", self)
            copy_text_act.triggered.connect(lambda: self._copy_tables_text(result))
            menu.addAction(copy_text_act)

            if _CHART_EDITOR_AVAILABLE and getattr(result, "figures", []):
                menu.addSeparator()
                edit_chart_act = QAction("✏️ 차트 수정", self)
                edit_chart_act.triggered.connect(self._edit_chart)
                menu.addAction(edit_chart_act)

            menu.addSeparator()

        copy_all_act = QAction("전체 복사", self)
        copy_all_act.triggered.connect(browser.selectAll)
        copy_all_act.triggered.connect(browser.copy)
        menu.addAction(copy_all_act)

        menu.exec(browser.viewport().mapToGlobal(pos))

    def _result_for_browser(self, browser: QTextBrowser):
        """현재 browser에 해당하는 AnalysisResult를 반환 (없으면 None)."""
        for i, r in enumerate(self._results):
            tab_idx = i + 1   # tab 0 = 로그
            if self.tab_widget.widget(tab_idx) is browser:
                return r
        return None

    def _copy_tables(self, result: Any) -> None:
        """결과 표를 text/html + text/plain(탭구분) 두 형식으로 클립보드에 복사."""
        from PySide6.QtCore import QMimeData

        tables = getattr(result, "tables", [])
        if not tables:
            QApplication.clipboard().setText(result.to_html())
            return

        html_parts = ["<html><body>"]
        plain_parts = []

        for tbl in tables:
            df = getattr(tbl, "dataframe", None)
            if df is None or len(df) == 0:
                continue
            title = getattr(tbl, "title", "")
            if title:
                html_parts.append(f"<p><strong>{title}</strong></p>")
                plain_parts.append(title)
            html_parts.append(df.to_html(index=False, border=1))
            plain_parts.append(df.to_csv(sep="\t", index=False))

        html_parts.append("</body></html>")
        full_html = "\n".join(html_parts)
        full_text = "\n".join(plain_parts)

        mime = QMimeData()
        mime.setHtml(full_html)
        mime.setText(full_text)
        QApplication.clipboard().setMimeData(mime)
        self.statusbar.showMessage("표가 클립보드에 복사되었습니다 (한글/Word에 붙여넣기 가능)")

    def _copy_tables_text(self, result: Any) -> None:
        """탭 구분 텍스트로만 복사."""
        tables = getattr(result, "tables", [])
        parts = []
        for tbl in tables:
            df = getattr(tbl, "dataframe", None)
            if df is not None and len(df) > 0:
                title = getattr(tbl, "title", "")
                if title:
                    parts.append(title)
                parts.append(df.to_csv(sep="\t", index=False))
        QApplication.clipboard().setText("\n".join(parts))
        self.statusbar.showMessage("텍스트가 클립보드에 복사되었습니다")

    # ------------------------------------------------------------------
    # 탭 닫기
    # ------------------------------------------------------------------

    def _on_tab_close(self, idx: int) -> None:
        if idx == 0:
            return  # 로그 탭은 닫기 불가
        result_idx = idx - 1
        if 0 <= result_idx < len(self._results):
            self._results.pop(result_idx)
        self.tab_widget.removeTab(idx)

    # ------------------------------------------------------------------
    # HTML 저장
    # ------------------------------------------------------------------

    def _save_html(self) -> None:
        path, _ = QFileDialog.getSaveFileName(
            self, "HTML 저장", "nuristat_output.html", "HTML (*.html)"
        )
        if not path:
            return
        try:
            browser = self.tab_widget.currentWidget()
            if isinstance(browser, QTextBrowser):
                html = browser.toHtml()
            else:
                html = "<html><body><p>내용 없음</p></body></html>"
            with open(path, "w", encoding="utf-8") as f:
                f.write(html)
            self.statusbar.showMessage(f"HTML 저장 완료: {path}")
        except Exception as exc:
            QMessageBox.critical(self, "오류", f"HTML 저장 실패:\n{exc}")

    # ------------------------------------------------------------------
    # Word 저장 (DataFrame 직접 빌드)
    # ------------------------------------------------------------------

    def _save_word(self) -> None:
        path, _ = QFileDialog.getSaveFileName(
            self, "Word 저장", "nuristat_output.docx", "Word (*.docx)"
        )
        if not path:
            return
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
        except ImportError:
            QMessageBox.critical(
                self, "오류",
                "python-docx 패키지가 필요합니다.\npip install python-docx"
            )
            return

        try:
            doc = Document()
            doc.core_properties.title = "누리스탯 분석 결과"
            heading = doc.add_heading("누리스탯 분석 결과", 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 현재 탭이 분석 결과 탭이면 그 결과만, 아니면 전체
            cur_idx = self.tab_widget.currentIndex()
            if cur_idx > 0 and cur_idx - 1 < len(self._results):
                results_to_export = [self._results[cur_idx - 1]]
            else:
                results_to_export = list(self._results)

            if not results_to_export:
                doc.add_paragraph("내보낼 분석 결과가 없습니다.")
                doc.save(path)
                self.statusbar.showMessage(f"Word 저장 완료: {path}")
                return

            for result in results_to_export:
                title = getattr(result, "title", "결과")
                doc.add_heading(title, level=1)

                # 표 직접 빌드
                for tbl in getattr(result, "tables", []):
                    df = getattr(tbl, "dataframe", None)
                    if df is None or len(df) == 0:
                        continue
                    tbl_title = getattr(tbl, "title", "")
                    if tbl_title:
                        doc.add_heading(tbl_title, level=3)

                    word_tbl = doc.add_table(rows=1, cols=len(df.columns))
                    word_tbl.style = "Table Grid"
                    hdr = word_tbl.rows[0].cells
                    for j, col in enumerate(df.columns):
                        hdr[j].text = str(col)
                    for _, row in df.iterrows():
                        cells = word_tbl.add_row().cells
                        for j, val in enumerate(row):
                            cells[j].text = str(val) if val is not None else ""
                    doc.add_paragraph()

                    for fn in getattr(tbl, "footnotes", []):
                        p = doc.add_paragraph(fn)
                        if p.runs:
                            p.runs[0].font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(9)

                # 이미지 (figures)
                for fig in getattr(result, "figures", []):
                    img_bytes = getattr(fig, "image_bytes", None)
                    if img_bytes:
                        try:
                            doc.add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
                            doc.add_paragraph()
                        except Exception:
                            pass

                # 주석 / 경고
                for note in getattr(result, "notes", []):
                    p = doc.add_paragraph(f"📌 {note}")
                for warn in getattr(result, "warnings", []):
                    p = doc.add_paragraph(f"⚠️ {warn}")

                doc.add_paragraph()   # 결과 간 간격

            doc.save(path)
            self.statusbar.showMessage(f"Word 저장 완료: {path}")
        except Exception as exc:
            QMessageBox.critical(self, "오류", f"Word 저장 실패:\n{exc}")

    # ------------------------------------------------------------------
    # 차트 편집
    # ------------------------------------------------------------------

    def _edit_chart(self) -> None:
        """현재 탭의 차트를 편집 대화상자로 수정하고 다시 렌더링합니다."""
        if not _CHART_EDITOR_AVAILABLE:
            QMessageBox.information(self, "안내", "차트 편집기를 불러올 수 없습니다.")
            return

        cur_idx = self.tab_widget.currentIndex()
        if cur_idx <= 0:
            QMessageBox.information(self, "안내", "차트 편집은 분석 결과 탭에서만 가능합니다.")
            return

        result_idx = cur_idx - 1
        if result_idx >= len(self._results):
            return
        result = self._results[result_idx]
        figures = getattr(result, "figures", [])
        if not figures:
            QMessageBox.information(self, "안내", "이 결과에 편집할 차트가 없습니다.")
            return

        try:
            from nuristat.analysis.visualization import VisualizationEngine
            _engine = VisualizationEngine()
        except Exception:
            QMessageBox.critical(self, "오류", "VisualizationEngine을 불러올 수 없습니다.")
            return

        # 첫 번째 Figure를 편집 (여러 개이면 이후 확장 가능)
        fig = figures[0]
        dlg = _ChartEditorDialog(fig, _engine.apply_edits, parent=self)
        if dlg.exec() == _ChartEditorDialog.Accepted:
            edited_fig = dlg.edited_figure()
            result.figures[0] = edited_fig
            # 탭 브라우저 재렌더링
            browser = self.tab_widget.widget(cur_idx)
            if isinstance(browser, QTextBrowser):
                ts = datetime.now().strftime("%H:%M:%S")
                title = getattr(result, "title", "결과")
                html = self._wrap_html(result.to_html(), title, ts)
                browser.setHtml(html)
            self.statusbar.showMessage("차트가 수정되었습니다.")

    # ------------------------------------------------------------------
    # 내부 헬퍼
    # ------------------------------------------------------------------

    def _wrap_html(self, body: str, title: str, ts: str) -> str:
        """body HTML을 테마 CSS + 헤더/타임스탬프로 감쌉니다."""
        styles = get_output_html_styles()
        return f"""<!DOCTYPE html>
<html>
<head>{styles}</head>
<body>
<h2>{title}</h2>
<div class="timestamp">{ts}</div>
{body}
</body>
</html>"""

    def closeEvent(self, event) -> None:
        self.hide()
        event.ignore()
